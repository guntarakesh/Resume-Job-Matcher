import os
import argparse
import pathlib
import tempfile
import logging
import json

import pdfplumber
import docx

from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import JsonOutputParser
from langchain_community.vectorstores import FAISS
from langchain_google_genai import (
    GoogleGenerativeAIEmbeddings,
    ChatGoogleGenerativeAI,
)
from dotenv import load_dotenv
load_dotenv()


try:
    import streamlit as st
    STREAMLIT = True
except ImportError:
    STREAMLIT = False

logging.getLogger("pdfminer").setLevel(logging.ERROR)

MATCHING_PROMPT_TEMPLATE = """
You are an experienced Technical Recruiter and ATS Resume Analyzer.

Compare the candidate's resume against the job description.

Perform the following analysis:

1. ATS Match Score (0-100)
2. Candidate Strengths
3. Missing Skills
4. Resume Improvement Suggestions
5. Recommended Interview Questions

Return ONLY valid JSON.

{{
    "score": 0,
    "strengths": [],
    "missing_skills": [],
    "resume_improvements": [],
    "interview_questions": []
}}

Resume:
{resume}

Job Description:
{job}

Retrieved Context:
{context}
"""

# Text extraction + caching
def extract_text(path: str) -> str:
    p = pathlib.Path(path)
    suffix = p.suffix.lower()
    try:
        if suffix == ".pdf":
            with pdfplumber.open(path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        if suffix in (".docx", ".doc"):
            doc = docx.Document(path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)
        return p.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        logging.warning("Failed to extract %s: %s", path, e)
        return ""

def clean_text(text: str) -> str:
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())

def extract_and_cache(path: str, cache_dir: str = "./cache") -> str:
    os.makedirs(cache_dir, exist_ok=True)
    src = pathlib.Path(path)
    cache_file = pathlib.Path(cache_dir) / (src.stem + ".txt")

    # Re-extract if cache missing or source modified after cache
    if cache_file.exists():
        try:
            cache_mtime = cache_file.stat().st_mtime
            src_mtime = src.stat().st_mtime
            if src_mtime <= cache_mtime:
                return cache_file.read_text(encoding="utf-8")
        except Exception:
            pass

    raw = extract_text(path)
    cleaned = clean_text(raw)
    if cleaned:
        try:
            cache_file.write_text(cleaned, encoding="utf-8")
        except Exception as e:
            logging.warning("Could not write cache %s: %s", cache_file, e)
    return cleaned

# Embeddings / Indexing
def get_embeddings():
    return GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-2",
        google_api_key=os.getenv("GOOGLE_API_KEY"),
    )

def index_documents(data_dir: str, index_dir: str):
    resumes = pathlib.Path(data_dir) / "resumes"
    jds = pathlib.Path(data_dir) / "jds"
    valid = {".pdf", ".txt", ".docx", ".doc"}
    docs = []

    for folder, source in ((resumes, "resume"), (jds, "job")):
        if not folder.exists():
            continue
        for f in folder.glob("**/*"):
            if f.suffix.lower() in valid:
                text = extract_and_cache(str(f))
                if text:
                    docs.append(Document(page_content=text, metadata={"source": source, "filename": f.name}))

    if not docs:
        print("[Warning] No documents to index.")
        return

    embeds = get_embeddings()
    if os.path.exists(index_dir):
        vs = FAISS.load_local(index_dir, embeds, allow_dangerous_deserialization=True, index_name="index_store")
        vs.add_documents(docs)
    else:
        vs = FAISS.from_documents(docs, embeds)

    os.makedirs(index_dir, exist_ok=True)
    vs.save_local(index_dir, "index_store")
    print(f"[Success] Indexed {len(docs)} documents.")

def get_retriever(index_dir: str, k: int = 5):
    if not os.path.exists(index_dir):
        print(f"[Error] Index directory not found: {index_dir}")
        return None
    embeds = get_embeddings()
    vs = FAISS.load_local(index_dir, embeds, allow_dangerous_deserialization=True, index_name="index_store")
    return vs.as_retriever(search_kwargs={"k": k})

# Matching
def match(resume_path: str, jd_path: str, index_dir: str, k: int = 5) -> dict:
    resume_text = extract_and_cache(resume_path)
    jd_text = extract_and_cache(jd_path)
    if not resume_text or not jd_text:
        return {"error": "Extraction failed for resume or JD."}

    retriever = get_retriever(index_dir, k)
    if not retriever:
        return {"error": "Could not load index."}

    llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    temperature=0,
    google_api_key=os.getenv("GOOGLE_API_KEY"),
    response_mime_type="application/json",
    )
    prompt = ChatPromptTemplate.from_template(MATCHING_PROMPT_TEMPLATE)

    rag_chain = (
        RunnablePassthrough.assign(
            context=(lambda x: x["resume"] + "\n" + x["job"]) | retriever,
        )
        | prompt
        | llm
        | JsonOutputParser()
    )

    try:
        return rag_chain.invoke({"resume": resume_text, "job": jd_text})
    except Exception as e:
        return {"error": f"LLM failed: {e}"}

# Streamlit UI
def run_ui(index_dir: str):
    with st.sidebar:
        st.header("About")

        st.write("""
    This application compares resumes against job descriptions using:

    - Gemini 2.5 Flash
    - LangChain
    - FAISS Vector Search
    - Retrieval-Augmented Generation (RAG)

    Developed by Rakesh Gunta.
    """)
    if not STREAMLIT:
        print("Streamlit not installed.")
        return
    st.title("🤖 AI Resume Intelligence Platform")
    st.caption("Powered by Gemini 2.5 Flash • LangChain • FAISS • Streamlit")
    r_file = st.file_uploader("📄 Upload Resume", type=["pdf", "txt", "docx", "doc"])
    j_file = st.file_uploader("📋 Upload Job Description", type=["pdf", "txt", "docx", "doc"])
    if r_file and j_file:
        temp_paths = {}
        for file, name in ((r_file, "r"), (j_file, "j")):
            suffix = pathlib.Path(file.name).suffix
            tf = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tf.write(file.read())
            tf.close()
            temp_paths[name] = tf.name

        if st.button("🚀 Analyze Resume"):
            with st.spinner("🤖 Gemini is analyzing your resume..."):
                res = match(temp_paths["r"], temp_paths["j"], index_dir)
                if "error" in res:
                    st.error(res["error"])
                elif "score" in res:
                    if res["score"] >= 85:
                        st.success("Excellent Match 🚀")

                    elif res["score"] >= 70:
                        st.info("Good Match 👍")

                    elif res["score"] >= 50:
                        st.warning("Average Match")

                    else:
                        st.error("Needs Improvement")

                    st.progress(min(max(res["score"] / 100, 0), 1))

                    st.divider()

                    st.subheader("✅ Matching Strengths")

                    for s in res.get("strengths", []):
                        st.success(s)

                    st.divider()

                    st.subheader("❌ Missing Skills")

                    for s in res.get("missing_skills", []):
                        st.warning(s)

                    st.divider()

                    st.subheader("💡 Resume Improvements")

                    for s in res.get("resume_improvements", []):
                        st.info(s)

                    st.divider()

                    st.subheader("🎤 Suggested Interview Questions")

                    for q in res.get("interview_questions", []):
                        st.write("•", q)

                    report = json.dumps(res, indent=4)

                    st.download_button(
                        label="📥 Download Report",
                        data=report,
                        file_name="resume_analysis.json",
                        mime="application/json",
                    )
                else:
                    st.write(res)

        for p in temp_paths.values():
            try:
                os.unlink(p)
            except Exception:
                pass

# CLI
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--action", choices=["index", "match", "serve"], default="serve")
    parser.add_argument("--data-dir", default="./data")
    parser.add_argument("--index-dir", default="./index")
    parser.add_argument("--resume")
    parser.add_argument("--jd")
    args = parser.parse_args()

    if args.action == "index":
        index_documents(args.data_dir, args.index_dir)
    elif args.action == "match":
        if not args.resume or not args.jd:
            print("--resume and --jd required")
            return
        print(match(args.resume, args.jd, args.index_dir))
    elif args.action == "serve":
        if not STREAMLIT:
            print("Install streamlit and run: streamlit run script.py -- --action serve")
            return
        run_ui(args.index_dir)

if __name__ == "__main__":
    main()
